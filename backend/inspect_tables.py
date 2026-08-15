import os
import sys
import docx
import json

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

def dump_file(path):
    print("=" * 60)
    print("FILE:", path)
    d = docx.Document(path)
    print("--- PARAGRAPHS ---")
    for i, p in enumerate(d.paragraphs[:12]):
        t = p.text.strip()
        if t:
            print(f"  P[{i}]: {t[:100]}")
    print("--- TABLES ---")
    for ti, t in enumerate(d.tables):
        print(f"  Table {ti} ({len(t.rows)} rows, {len(t.columns)} cols):")
        for ri, r in enumerate(t.rows[:8]):
            # deduplicate cell text
            row_vals = []
            seen = set()
            for c in r.cells:
                ct = c.text.strip().replace('\n', ' ')
                row_vals.append(ct)
            print(f"    R[{ri}]: {row_vals[:4]}")

if __name__ == '__main__':
    # Sample files from Stomatology & Davolash Ishi
    sample_files = [
        'datas/STOMOTOLOGY/MODULE 1/Dental_Pain_Sensitivity_Smart_Phrasebook_with_Patient_Language -1 MODULE.docx',
        'datas/STOMOTOLOGY/MODULE 1/Dental_Pain_Sensitivity_1 MODULE_Grammar.docx',
        'datas/STOMOTOLOGY/MODULE 1/Dental_Pain_Sensitivity_Natural_Dialogue-1 MODULE.docx',
        'datas/DAVOLASH ISHI/Module_01/Module_01_Vocabulary_Smart_Phrasebook.docx',
        'datas/DAVOLASH ISHI/Module_01/Module_01_Grammar_Present_Simple_Present_Continuous.docx',
        'datas/DAVOLASH ISHI/Module_01/01_Dialogue_B2.docx',
    ]
    for sf in sample_files:
        if os.path.exists(sf):
            dump_file(sf)
