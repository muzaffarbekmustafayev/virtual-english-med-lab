import os
import json
import docx

def read_docx(path):
    doc = docx.Document(path)
    lines = []
    # Read paragraphs
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    # Read tables
    tables_data = []
    for table in doc.tables:
        t_rows = []
        for row in table.rows:
            r_cells = [cell.text.strip() for cell in row.cells]
            if any(r_cells):
                t_rows.append(r_cells)
        if t_rows:
            tables_data.append(t_rows)
    return lines, tables_data

def inspect_all():
    summary = {}
    for root, dirs, files in os.walk('datas'):
        for f in files:
            if f.endswith('.docx') and not f.startswith('~$'):
                p = os.path.join(root, f)
                try:
                    lines, tables = read_docx(p)
                    summary[p] = {
                        'lines_count': len(lines),
                        'sample_lines': lines[:10],
                        'tables_count': len(tables),
                        'sample_table': tables[0][:4] if tables else None
                    }
                except Exception as e:
                    summary[p] = {'error': str(e)}

    print(f"Total DOCX files parsed: {len(summary)}")
    with open('datas_summary.json', 'w', encoding='utf-8') as out:
        json.dump(summary, out, indent=2, ensure_ascii=False)
    print("Saved datas_summary.json")

if __name__ == '__main__':
    inspect_all()
