import os
import sys
import json
import re
import docx

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

def clean(text):
    if not text:
        return ""
    t = text.replace('\r\n', '\n').replace('\r', '\n')
    t = re.sub(r'[ \t]+', ' ', t)
    t = t.replace('‘', "'").replace('’', "'").replace('“', '"').replace('”', '"')
    t = t.replace('–', '-').replace('—', '-')
    return t.strip()

def read_docx_clean(path):
    doc = docx.Document(path)
    paragraphs = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    tables = []
    for table in doc.tables:
        t_rows = []
        for row in table.rows:
            row_cells = [clean(cell.text) for cell in row.cells]
            if any(row_cells):
                t_rows.append(row_cells)
        if t_rows:
            tables.append(t_rows)
    return paragraphs, tables

def parse_dialogue_file(paragraphs):
    turns = []
    current_speaker = None
    current_text = []

    for p in paragraphs:
        m = re.match(r'^(Doctor|Patient|Shifokor|Bemor|Dr\.|Pt\.)\s*:\s*(.*)', p, re.IGNORECASE)
        if m:
            if current_speaker and current_text:
                turns.append({'speaker': current_speaker, 'text': clean(' '.join(current_text))})
            speaker_raw = m.group(1).lower()
            current_speaker = 'Doctor' if any(k in speaker_raw for k in ['doc', 'shifokor', 'dr']) else 'Patient'
            current_text = [m.group(2)]
        elif current_speaker:
            current_text.append(p)
        else:
            if ':' in p:
                parts = p.split(':', 1)
                spk = parts[0].strip().lower()
                if any(k in spk for k in ['doc', 'shifokor', 'dr']):
                    current_speaker = 'Doctor'
                    current_text = [parts[1].strip()]
                elif any(k in spk for k in ['pat', 'bemor', 'pt']):
                    current_speaker = 'Patient'
                    current_text = [parts[1].strip()]

    if current_speaker and current_text:
        turns.append({'speaker': current_speaker, 'text': clean(' '.join(current_text))})
    return turns

def make_clinical_definitions(word, uz, ru, raw_def=""):
    w_lower = word.lower()
    
    en_def = f"Clinical medical condition and symptom associated with {word.lower()}."
    uz_def = f"{uz or word} - bemor ko'rigida aniqlanadigan klinik belgi yoki holat."
    ru_def = f"{ru or word} - клинический симптом или состояние, выявляемое при осмотре."

    if 'pain' in w_lower or 'ache' in w_lower:
        en_def = f"An unpleasant sensory and emotional experience in the oral or anatomical region representing {word.lower()}."
        uz_def = f"{uz or word} - tish yoki to'qimalar sohasidagi yoqimsiz og'riq va noqulaylik hissi."
        ru_def = f"{ru or word} - неприятное болевое ощущение или дискомфорт в тканях."
    elif 'sensitivity' in w_lower or 'sensitive' in w_lower:
        en_def = "Discomfort or pain in teeth triggered by thermal (hot/cold), sweet, or tactile stimuli."
        uz_def = "Tishning issiq, sovuq, shirin yoki mexanik ta'sirlarga nisbatan o'ta sezgirligi."
        ru_def = "Повышенная чувствительность зубов к термическим или химическим раздражителям."
    elif 'swelling' in w_lower or 'abscess' in w_lower or 'edema' in w_lower:
        en_def = "Localized enlargement or accumulation of inflammatory fluid and pus in tissue spaces."
        uz_def = "To'qimalarda yallig'lanish, suyuqlik yoki yiring to'planishi tufayli yuzaga kelgan shish."
        ru_def = "Воспалительное увеличение или скопление гнойного экссудата в тканях."
    elif 'bleed' in w_lower or 'hemorrhage' in w_lower:
        en_def = "Escape of blood from damaged blood vessels or inflamed gingival margins."
        uz_def = "Shikastlangan qon tomirlari yoki yallig'langan milkdan qon ketishi."
        ru_def = "Кровоточивость из поврежденных сосудов или воспаленных десен."
    elif 'caries' in w_lower or 'decay' in w_lower or 'cavity' in w_lower:
        en_def = "Bacterial demineralization and progressive destruction of hard tooth structure."
        uz_def = "Bakteriyalar ta'sirida tish qattiq to'qimalarining emirilishi va kovak hosil bo'lishi."
        ru_def = "Бактериальное разрушение твердых тканей зуба с образованием кариозной полости."
    elif 'extract' in w_lower or 'removal' in w_lower:
        en_def = "Surgical extraction of a non-restorable tooth from the alveolar bone socket."
        uz_def = "Davolash imkoni bo'lmagan tishni jag' suyagi katakchasidan jarrohlik yo'li bilan sug'urib olish."
        ru_def = "Хирургическое удаление зуба из костной альвеолы."
    elif 'gingiv' in w_lower or 'gum' in w_lower or 'periodont' in w_lower:
        en_def = "Inflammatory disease affecting the gingival tissues and periodontal support of teeth."
        uz_def = "Milk to'qimalari va tishni ushlab turuvchi parodont to'qimalarining yallig'lanishi."
        ru_def = "Воспалительное поражение десен и удерживающих тканей пародонта."
    elif 'fever' in w_lower or 'temperature' in w_lower:
        en_def = "Systemic elevation of body core temperature in response to an infectious process."
        uz_def = "Infeksion yallig'lanish jarayoniga javoban tana haroratining ko'tarilishi."
        ru_def = "Повышение температуры тела в ответ на инфекционный процесс."
    elif 'dyspnea' in w_lower or 'breath' in w_lower:
        en_def = "Subjective clinical sensation of difficult, labored, or uncomfortable breathing."
        uz_def = "Nafas olishning qiyinlashishi, havo yetishmasligi hissi (nafas qisishi)."
        ru_def = "Субъективное ощущение нехватки воздуха и затруднения дыхания (одышка)."
    elif 'chest' in w_lower or 'angina' in w_lower:
        en_def = "Discomfort or constricting pressure localized in the thoracic region."
        uz_def = "Ko'krak qafasi sohasida siquvchi yoki bosuvchi og'riq hissi."
        ru_def = "Сдавливающая или давящая боль в области грудной клетки."

    return {
        'en': en_def,
        'uz': uz_def,
        'ru': ru_def
    }

def parse_vocab_tables(tables, paragraphs):
    vocab_list = []
    seen_words = set()

    for t in tables:
        if not t or len(t) < 2:
            continue
        header = [clean(c).lower() for c in t[0]]
        
        # Check if table is vocabulary table
        is_vocab = any(k in ' '.join(header) for k in ['term', 'word', 'so‘z', 'atama', 'lug‘at', 'pronunciation', 'talaffuz', 'english', 'русский', 'o‘zbekcha'])
        is_phrase = any(k in ' '.join(header) for k in ['doctor', 'phrase', 'ibora', 'bosqich', 'savol', 'gap'])

        if not is_vocab and is_phrase:
            continue

        # Find columns
        word_col = 0
        uz_col = -1
        ru_col = -1
        def_col = -1
        ex_col = -1

        for ci, c in enumerate(header):
            if any(k in c for k in ['english', 'term', 'word', 'atama', 'so‘z']):
                word_col = ci
            elif any(k in c for k in ['o‘zbek', "o'zbek", 'uzbek', 'tarjima', 'o‘zbekcha']):
                uz_col = ci
            elif any(k in c for k in ['rus', 'рус', 'russian', 'русский']):
                ru_col = ci
            elif any(k in c for k in ['def', "ta'rif", 'maʼno', 'значение']):
                def_col = ci
            elif any(k in c for k in ['ex', 'misol', 'namuna', 'пример']):
                ex_col = ci

        # If header had '#' column at 0
        if '#' in header[0] and len(header) > 1:
            word_col = 1
            if uz_col == 1:
                uz_col = 2 if len(header) > 2 else 1

        if uz_col == -1 and len(header) >= 3:
            uz_col = 2
        if ru_col == -1 and len(header) >= 4:
            ru_col = 3

        for row in t[1:]:
            if len(row) <= word_col:
                continue
            w = clean(row[word_col])
            # strip numeric prefix "1.", "1)"
            w = re.sub(r'^\d+[\.\)]\s*', '', w)
            if not w or len(w) < 2 or w.lower() in ['#', 'term', 'word', 'atamalar', 'so‘z', 'english']:
                continue
            if w.lower() in seen_words:
                continue

            uz = clean(row[uz_col]) if uz_col != -1 and len(row) > uz_col else ""
            ru = clean(row[ru_col]) if ru_col != -1 and len(row) > ru_col else ""
            df = clean(row[def_col]) if def_col != -1 and len(row) > def_col else ""
            ex = clean(row[ex_col]) if ex_col != -1 and len(row) > ex_col else ""

            defs = make_clinical_definitions(w, uz, ru, df)

            seen_words.add(w.lower())
            vocab_list.append({
                'word': w,
                'translation': uz or w,
                'translation_uz': uz or w,
                'translation_ru': ru or "",
                'translation_en': w,
                'definition': defs['en'],
                'definition_uz': defs['uz'],
                'definition_ru': defs['ru'],
                'definition_en': defs['en'],
                'example': ex or f"The doctor assessed the patient for {w.lower()}."
            })

    # Also parse from bullet points or lines if table had few entries
    if len(vocab_list) < 5:
        for p in paragraphs:
            if ' - ' in p or ' — ' in p:
                parts = re.split(r'\s*[-—]\s*', p, maxsplit=1)
                if len(parts) == 2 and len(parts[0]) < 40 and not parts[0].startswith('Doctor') and not parts[0].startswith('Patient'):
                    w = re.sub(r'^\d+[\.\)]\s*', '', parts[0]).strip()
                    if w and w.lower() not in seen_words and len(w) > 2:
                        uz_val = parts[1].strip()
                        defs = make_clinical_definitions(w, uz_val, "")
                        seen_words.add(w.lower())
                        vocab_list.append({
                            'word': w,
                            'translation': uz_val,
                            'translation_uz': uz_val,
                            'translation_ru': "",
                            'translation_en': w,
                            'definition': defs['en'],
                            'definition_uz': defs['uz'],
                            'definition_ru': defs['ru'],
                            'definition_en': defs['en'],
                            'example': f"The doctor noted {w.lower()} during the exam."
                        })

    return vocab_list

def parse_phrasebook_tables(tables, paragraphs):
    phrases = []
    seen_phrases = set()

    for t in tables:
        if not t or len(t) < 2:
            continue
        header = [c.lower() for c in t[0]]
        
        # Check if table has phrases
        is_phrase_table = any(k in ' '.join(header) for k in ['doctor', 'phrase', 'ibora', 'bosqich', 'patient', 'savol'])
        if not is_phrase_table and not any('?' in str(cell) for row in t for cell in row):
            continue

        cat_col = -1
        doc_col = 0
        uz_col = -1
        ru_col = -1
        pat_col = -1

        for ci, c in enumerate(header):
            if any(k in c for k in ['bosqich', 'kategoriya', 'category', 'step', 'stage']):
                cat_col = ci
            elif any(k in c for k in ['doctor', 'phrase', 'shifokor', 'savol', 'ibora']):
                doc_col = ci
            elif any(k in c for k in ['patient', 'bemor']):
                pat_col = ci
            elif any(k in c for k in ['o‘zbek', "o'zbek", 'uzbek', 'tarjima']):
                if uz_col == -1:
                    uz_col = ci
            elif 'rus' in c:
                if ru_col == -1:
                    ru_col = ci

        if doc_col == 0 and cat_col == 0 and len(header) > 1:
            doc_col = 1
        if uz_col == -1 and len(header) > 2:
            uz_col = 2
        if ru_col == -1 and len(header) > 3:
            ru_col = 3

def get_clinical_phrase_translations(phrase, uz_raw="", ru_raw=""):
    p_lower = phrase.lower()
    uz = ""
    ru = ""

    if 'what brings you in' in p_lower:
        uz = "Bugun sizni nima bezovta qilib, olib keldi?"
        ru = "Что привело вас к нам сегодня? На что жалуетесь?"
    elif 'tell me about your' in p_lower or 'tell me what' in p_lower:
        uz = "Tish / shikoyatingiz haqida batafsil aytib bera olasizmi?"
        ru = "Расскажите подробнее о вашей боли / проблеме?"
    elif 'how long have you had' in p_lower or 'how long has' in p_lower or 'how long have' in p_lower:
        uz = "Bu og'riq / simptom qachondan beri bor?"
        ru = "Как долго вас беспокоит эта боль / симптом?"
    elif 'when did the pain start' in p_lower or 'when did it start' in p_lower:
        uz = "Og'riq qachon boshlandi?"
        ru = "Когда именно началась боль?"
    elif 'become worse' in p_lower or 'getting worse' in p_lower:
        uz = "Og'riq kuchaydimi yoki o'zgarmadimi?"
        ru = "Боль усилилась или осталась прежней?"
    elif 'cold or hot' in p_lower or 'hot or cold' in p_lower:
        uz = "Sovuq yoki issiq narsa yeb-ichganingizda og'riydimi?"
        ru = "Болит ли при употреблении холодного или горячего?"
    elif 'what usually makes the pain worse' in p_lower or 'makes it worse' in p_lower:
        uz = "Odatda og'riqni nima kuchaytiradi?"
        ru = "Что обычно провоцирует или усиливает боль?"
    elif 'how would you describe' in p_lower:
        uz = "Og'riqni qanday ta'riflaysiz (o'tkir, simillovchi, urib turuvchi)?"
        ru = "Как бы вы описали боль (острая, ноющая, пульсирующая)?"
    elif 'is the pain constant' in p_lower or 'constant or' in p_lower:
        uz = "Og'riq doimiymi yoki kelib-ketib turadimi?"
        ru = "Боль постоянная или периодическая?"
    elif 'scale of 0 to 10' in p_lower or 'rate the' in p_lower or '0-10' in p_lower:
        uz = "Eng kuchli og'riqni 0 dan 10 gacha necha ballga baholaysiz?"
        ru = "Оцените интенсивность боли по шкале от 0 до 10?"
    elif 'how severe' in p_lower:
        uz = "Og'riq qanchalik kuchli?"
        ru = "Насколько сильна эта боль?"
    elif 'where exactly' in p_lower or 'where is the' in p_lower:
        uz = "Og'riqni aynan qayerda his qilasiz?"
        ru = "Где именно вы ощущаете боль?"
    elif 'spread' in p_lower or 'radiat' in p_lower:
        uz = "Og'riq boshqa joyga (quloq, jag', bo'yin) tarqaladimi?"
        ru = "Отдает (иррадиирует) ли боль в челюсть, ухо или шею?"
    elif 'without eating' in p_lower or 'spontaneous' in p_lower:
        uz = "Hech narsa yemay-ichmay turganda ham o'z-o'zidan og'riq bo'ladimi?"
        ru = "Возникает ли самопроизвольная боль в покое, без еды?"
    elif 'how long does the pain usually last' in p_lower or 'how long does it last' in p_lower:
        uz = "Qo'zg'atuvchidan keyin og'riq qancha vaqt davom etadi?"
        ru = "Как долго длится боль после воздействия раздражителя?"
    elif 'discomfort during the examination' in p_lower or 'let me know' in p_lower:
        uz = "Tekshiruv vaqtida noqulaylik sezsangiz, darhol menga ayting."
        ru = "Пожалуйста, дайте знать, если почувствуете дискомфорт при осмотре."
    elif 'examine the tooth' in p_lower or 'examine you' in p_lower or 'take a look' in p_lower:
        uz = "Endi tishni / sohani tekshirib ko'rmoqchiman."
        ru = "Сейчас я проведу осмотр беспокоящего зуба / области."
    elif 'any other symptoms' in p_lower:
        uz = "Boshqa qo'shimcha simptomlar yoki bezovtaliklar bormi?"
        ru = "Есть ли другие симптомы или жалобы?"
    elif 'take any medication' in p_lower or 'painkillers' in p_lower:
        uz = "Og'riq qoldiruvchi dori qabul qildingizmi?"
        ru = "Принимали ли вы обезболивающие препараты?"
    elif 'bleeding' in p_lower:
        uz = "Qonash kuzatiladimi yoki milk shishganmi?"
        ru = "Наблюдается ли кровоточивость или припухлость десен?"
    elif 'swelling' in p_lower:
        uz = "Shish yoki yiringli ajralma bormi?"
        ru = "Есть ли припухлость или гнойные выделения?"
    elif 'fever' in p_lower or 'temperature' in p_lower:
        uz = "Harorat (isitma) yoki qaltirash bormi?"
        ru = "Есть ли повышенная температура или озноб?"

    # Fallback to provided raw or default
    uz = uz or uz_raw or "Klinik muloqot iborasi"
    ru = ru or (ru_raw if ru_raw and ru_raw != uz_raw else "")
    if not ru:
        ru = "Клиническая фраза для сбора анамнеза"

    return {
        'uz': uz,
        'ru': ru
    }

def parse_phrasebook_tables(tables, paragraphs):
    phrases = []
    seen_phrases = set()

    for t in tables:
        if not t or len(t) < 2:
            continue
        header = [c.lower() for c in t[0]]
        
        # Check if table has phrases
        is_phrase_table = any(k in ' '.join(header) for k in ['doctor', 'phrase', 'ibora', 'bosqich', 'patient', 'savol'])
        if not is_phrase_table and not any('?' in str(cell) for row in t for cell in row):
            continue

        cat_col = -1
        doc_col = 0
        uz_col = -1
        ru_col = -1
        pat_col = -1

        for ci, c in enumerate(header):
            if any(k in c for k in ['bosqich', 'kategoriya', 'category', 'step', 'stage']):
                cat_col = ci
            elif any(k in c for k in ['doctor', 'phrase', 'shifokor', 'savol', 'ibora']):
                doc_col = ci
            elif any(k in c for k in ['patient', 'bemor']):
                pat_col = ci
            elif any(k in c for k in ['o‘zbek', "o'zbek", 'uzbek', 'tarjima']):
                if uz_col == -1:
                    uz_col = ci
            elif 'rus' in c:
                if ru_col == -1:
                    ru_col = ci

        if doc_col == 0 and cat_col == 0 and len(header) > 1:
            doc_col = 1
        if uz_col == -1 and len(header) > 2:
            uz_col = 2
        if ru_col == -1 and len(header) > 3:
            ru_col = 3

        for row in t[1:]:
            if len(row) <= doc_col:
                continue
            p_text = clean(row[doc_col])
            p_text = re.sub(r'^\d+[\.\)]\s*', '', p_text)
            if not p_text or len(p_text) < 4 or p_text.lower() in ['doctor', 'phrase', 'ibora', 'bosqich']:
                continue
            if p_text.lower() in seen_phrases:
                continue

            category = clean(row[cat_col]) if cat_col != -1 and len(row) > cat_col else "Clinical Consultation"
            if not category or category.isdigit():
                category = "Clinical Consultation"

            uz = row[uz_col] if uz_col != -1 and len(row) > uz_col else ""
            ru = row[ru_col] if ru_col != -1 and len(row) > ru_col else ""
            pat = row[pat_col] if pat_col != -1 and len(row) > pat_col else ""

            t_hints = get_clinical_phrase_translations(p_text, uz, ru)

            seen_phrases.add(p_text.lower())
            phrases.append({
                'phrase': p_text,
                'category': category,
                'hint': t_hints['uz'],
                'hint_uz': t_hints['uz'],
                'hint_ru': t_hints['ru'],
                'hint_en': "Clinical consultation inquiry",
                'patient_response': pat
            })

    # Also parse from paragraphs if structured
    current_cat = "Clinical Inquiry"
    for p in paragraphs:
        if any(cat.lower() in p.lower() for cat in ['Opening', 'Onset', 'Duration', 'Trigger', 'Character', 'Severity', 'Location', 'Clarification', 'Treatment', 'Summary', 'Aftercare', 'Examination']):
            for cat_name in ['Opening', 'Onset & Duration', 'Triggers', 'Pain Character', 'Severity', 'Location', 'Clarification', 'Treatment Plan', 'Aftercare', 'Examination', 'Summary']:
                if cat_name.lower() in p.lower():
                    current_cat = cat_name
                    break
        elif '?' in p and any(w in p.lower() for w in ['what', 'how', 'when', 'where', 'do you', 'does', 'have you', 'can you', 'could you', 'are you']):
            lines = [l.strip() for l in p.split('\n') if l.strip()]
            phr_text = lines[0]
            phr_text = re.sub(r'^\d+[\.\)]\s*', '', phr_text)
            if phr_text.lower() not in seen_phrases and len(phr_text) > 5:
                uz_hint = ""
                ru_hint = ""
                for l in lines[1:]:
                    if any(k in l.lower() for k in ['uzbek:', 'uzb:', "o'zbek:"]):
                        uz_hint = re.sub(r'^(uzbek|uzb|o‘zbek|o\'zbek):\s*', '', l, flags=re.IGNORECASE).strip()
                    elif any(k in l.lower() for k in ['russian:', 'rus:']):
                        ru_hint = re.sub(r'^(russian|rus):\s*', '', l, flags=re.IGNORECASE).strip()

                t_hints = get_clinical_phrase_translations(phr_text, uz_hint, ru_hint)

                seen_phrases.add(phr_text.lower())
                phrases.append({
                    'phrase': phr_text,
                    'category': current_cat,
                    'hint': t_hints['uz'],
                    'hint_uz': t_hints['uz'],
                    'hint_ru': t_hints['ru'],
                    'hint_en': "Clinical consultation inquiry",
                    'patient_response': ""
                })

    return phrases

def parse_grammar_details(paragraphs, tables):
    grammar_items = []
    title = "Clinical Grammar Focus"
    rules_uz = []
    rules_ru = []
    rules_en = []
    pattern = ""
    examples = []

    # Title extraction
    for p in paragraphs[:6]:
        if any(k in p.lower() for k in ['grammar', 'present', 'past', 'future', 'modal', 'passive', 'conditional', 'imperative']):
            clean_title = re.sub(r'^(module\s*\d+\s*[-—:]*|\d+[\.\)]*)\s*', '', p, flags=re.IGNORECASE).strip()
            if clean_title:
                title = clean_title
                break

    # Paragraph rules and dialog examples extraction
    for i, p in enumerate(paragraphs):
        p_clean = clean(p)
        if len(p_clean) > 20 and not p_clean.startswith('Doctor:') and not p_clean.startswith('Patient:'):
            if any(k in p_clean.lower() for k in ['qoida', 'ishlatiladi', 'bemor', 'foydalaniladi', 'talaba']):
                rules_uz.append(p_clean)
            elif any(k in p_clean.lower() for k in ['правило', 'используется', 'студент', 'пациент']):
                rules_ru.append(p_clean)
            elif any(k in p_clean.lower() for k in ['rule', 'used', 'clinician', 'when', 'structure']):
                rules_en.append(p_clean)
            elif any(k in p_clean.lower() for k in ['formula', 'structure:', 'pattern:']) or '+' in p_clean:
                if not pattern and len(p_clean) < 120:
                    pattern = p_clean

        # Check for dialogue quotes like “Cold drinks usually make it worse.”
        quotes = re.findall(r'[“\"\'](.*?[a-zA-Z]{3,}.*?)[”\"\']', p_clean)
        if quotes and any(k in p_clean.lower() for k in ['misol', 'example', 'dialog', 'пример']):
            sent_eng = quotes[0].strip()
            if len(sent_eng) > 6 and not sent_eng.lower().startswith('http') and not any(x in sent_eng for x in ['qoida', 'formula', 'zamoni']):
                uz_t = ""
                ru_t = ""
                if i + 1 < len(paragraphs):
                    cand1 = clean(paragraphs[i + 1])
                    if not cand1.startswith('“') and not cand1.startswith('"') and len(cand1) < 160:
                        uz_t = cand1
                if i + 2 < len(paragraphs):
                    cand2 = clean(paragraphs[i + 2])
                    if not cand2.startswith('“') and not cand2.startswith('"') and len(cand2) < 160:
                        ru_t = cand2

                if not any(ex['sentence'] == sent_eng for ex in examples):
                    examples.append({
                        'sentence': sent_eng,
                        'translation': uz_t or sent_eng,
                        'translation_uz': uz_t or sent_eng,
                        'translation_ru': ru_t or "",
                        'note': 'Clinical Dialogue Example'
                    })

    # Table examples extraction
    if tables:
        for t in tables:
            if not t or len(t) < 2:
                continue
            header = [clean(c).lower() for c in t[0]]
            eng_col = 0
            uz_col = -1
            ru_col = -1
            note_col = -1

            for ci, c in enumerate(header):
                if any(k in c for k in ['signal', 'formula', 'english', 'example', 'misol', 'gap']):
                    eng_col = ci
                elif any(k in c for k in ['o‘zbek', "o'zbek", 'uzbek', 'tarjima']):
                    uz_col = ci
                elif any(k in c for k in ['rus', 'рус', 'russian']):
                    ru_col = ci
                elif any(k in c for k in ['zamon', 'tense', 'note', 'izoh']):
                    note_col = ci

            if uz_col == -1 and len(header) >= 3:
                uz_col = 2
            if ru_col == -1 and len(header) >= 4:
                ru_col = 3
            if note_col == -1 and len(header) >= 2:
                note_col = 1

            for r in t[1:]:
                if len(r) <= eng_col:
                    continue
                sent = clean(r[eng_col])
                if not sent or len(sent) < 3 or sent.lower() in ['#', 'signal / formula', 'english', 'gap']:
                    continue

                uz_val = clean(r[uz_col]) if uz_col != -1 and len(r) > uz_col else ""
                ru_val = clean(r[ru_col]) if ru_col != -1 and len(r) > ru_col else ""
                note_val = clean(r[note_col]) if note_col != -1 and len(r) > note_col else "Clinical Pattern"

                if not any(ex['sentence'] == sent for ex in examples):
                    examples.append({
                        'sentence': sent,
                        'translation': uz_val or sent,
                        'translation_uz': uz_val or sent,
                        'translation_ru': ru_val or "",
                        'note': note_val
                    })

    if not pattern:
        pattern = "Subject + Auxiliary Verb + Main Verb (Clinical Context)"

    rule_text_uz = ' '.join(rules_uz[:2]) if rules_uz else "Klinik muloqotda aniq grammatik qoidalar bemor holatini to'g'ri baholashga yordam beradi."
    rule_text_ru = ' '.join(rules_ru[:2]) if rules_ru else "В клиническом диалоге точные грамматические структуры помогают быстрее понять жалобы пациента."
    rule_text_en = ' '.join(rules_en[:2]) if rules_en else "Using precise clinical grammatical structures enables clear patient history taking and assessment."

    grammar_items.append({
        'title': title,
        'title_uz': title,
        'title_ru': title,
        'title_en': title,
        'rule_explanation': rule_text_uz,
        'rule_explanation_uz': rule_text_uz,
        'rule_explanation_ru': rule_text_ru,
        'rule_explanation_en': rule_text_en,
        'structure_pattern': pattern,
        'examples': examples[:8] if examples else [
            {
                'sentence': "How long has this symptom been bothering you?",
                'translation': "Bu simptom sizni qancha vaqtdan beri bezovta qilmoqda?",
                'translation_uz': "Bu simptom sizni qancha vaqtdan beri bezovta qilmoqda?",
                'translation_ru': "Как долго вас беспокоит этот симптом?",
                'note': "Clinical duration inquiry"
            },
            {
                'sentence': "Does the pain get worse with hot or cold temperatures?",
                'translation': "Og'riq issiq yoki sovuq haroratda kuchayadimi?",
                'translation_uz': "Og'riq issiq yoki sovuq haroratda kuchayadimi?",
                'translation_ru': "Усиливается ли боль от горячего или холодного?",
                'note': "Clinical trigger inquiry"
            }
        ],
        'common_mistakes': [
            {
                'incorrect': "Since when you feel this pain?",
                'correct': "How long have you had this pain?",
                'explanation': "Davomiylikni so'rashda 'How long have you had...?' standarti ishlatiladi.",
                'explanation_uz': "Davomiylikni so'rashda 'How long have you had...?' standarti ishlatiladi.",
                'explanation_ru': "Для выяснения длительности используется конструкция 'How long have you had...?'.",
                'explanation_en': "Use 'How long have you had...?' when asking about the duration of symptoms."
            },
            {
                'incorrect': "Where is paining you?",
                'correct': "Where exactly do you feel the pain?",
                'explanation': "Lokatsiyani aniqlashda 'Where do you feel the pain?' to'g'ri grammatik shakl.",
                'explanation_uz': "Lokatsiyani aniqlashda 'Where do you feel the pain?' to'g'ri grammatik shakl.",
                'explanation_ru': "Правильный вопрос о локализации боли: 'Where exactly do you feel the pain?'.",
                'explanation_en': "Ask 'Where exactly do you feel the pain?' to identify the specific anatomical site."
            }
        ]
    })
    return grammar_items

def generate_module_quizzes(module_title, vocab_list, grammar_items):
    quizzes = []
    if len(vocab_list) >= 4:
        target = vocab_list[0]
        others = [v['translation'] for v in vocab_list[1:4]]
        quizzes.append({
            'question': f"What is the medical meaning of the term '{target['word']}' in clinical consultation?",
            'option_a': target['translation'],
            'option_b': others[0],
            'option_c': others[1],
            'option_d': others[2],
            'correct_option': 'A',
            'explanation': f"'{target['word']}' atamasining to'g'ri ma'nosi: {target['translation']}."
        })

    if len(vocab_list) >= 8:
        target2 = vocab_list[4]
        quizzes.append({
            'question': f"Which term corresponds to '{target2['translation']}'?",
            'option_a': vocab_list[1]['word'],
            'option_b': target2['word'],
            'option_c': vocab_list[2]['word'],
            'option_d': vocab_list[3]['word'],
            'correct_option': 'B',
            'explanation': f"'{target2['translation']}' ma'nosiga mos keluvchi tibbiy termin '{target2['word']}' hisoblanadi."
        })

    quizzes.append({
        'question': "Which of the following is the grammatically and clinically standard question to ask about the duration of symptoms?",
        'option_a': "How long do you having this symptom?",
        'option_b': "How long has this condition been bothering you?",
        'option_c': "Since what time you have this problem?",
        'option_d': "How many days you are feeling pain?",
        'correct_option': 'B',
        'explanation': "Klinik anamnezda davomiylikni so'rash uchun Present Perfect Continuous ('How long has... been bothering you?') to'g'ri shakl hisoblanadi."
    })

    quizzes.append({
        'question': "When evaluating acute pain, what is the best phrasing to assess pain character?",
        'option_a': "Is your pain bad or very bad?",
        'option_b': "Can you describe the pain—is it sharp, dull, throbbing, or aching?",
        'option_c': "Why do you have sharp pain today?",
        'option_d': "Stop pain with pills immediately.",
        'correct_option': 'B',
        'explanation': "Og'riq xarakterini aniqlashda shifokor ochiq va taklif qiluvchi variantlar (sharp, dull, throbbing, aching) beradi."
    })

    return quizzes

def build_curriculum():
    print("=== Extracting & Formatting All 20 Medical Modules ===")

    specs = [
        {
            'id': 1,
            'name': 'Stomatologiya',
            'name_uz': 'Stomatologiya',
            'name_ru': 'Стоматология',
            'name_en': 'Dentistry',
            'folder': 'datas/STOMOTOLOGY'
        },
        {
            'id': 2,
            'name': 'Davolash ishi',
            'name_uz': 'Davolash ishi',
            'name_ru': 'Лечебное дело',
            'name_en': 'General Medicine',
            'folder': 'datas/DAVOLASH ISHI'
        }
    ]

    all_modules = []

    titles_map = {
        (1, 1): ("Dental Pain & Sensitivity", "Tish og'rig'i va sezgirligi", "Зубная боль и чувствительность"),
        (1, 2): ("Tooth Extraction Consultation", "Tishni sug'urish bo'yicha maslahat", "Консультация по удалению зуба"),
        (1, 3): ("Acute Toothache & Pulpitis", "O'tkir tish og'rig'i va pulpit", "Острая зубная боль и пульпит"),
        (1, 4): ("Dental Abscess & Swelling", "Dental abssess va shish", "Дентальный абсцесс и отек"),
        (1, 5): ("Dental Caries & Restorations", "Tish kariyesi va plombalash", "Кариес зубов и реставрация"),
        (1, 6): ("Gum Problems & Periodontitis", "Milk kasalliklari va parodontit", "Заболевания десен и пародонтит"),
        (1, 7): ("Impacted Wisdom Tooth", "Retensiyalangan aql tishi", "Ретинированный зуб мудрости"),
        (1, 8): ("Dental Emergency & Trauma", "Shoshilinch stomatologik jarohat", "Неотложная стоматологическая травма"),
        (1, 9): ("Prosthodontics & Crown Restoration", "Protezlash va tish qoplamalari", "Протезирование и коронки"),
        (1, 10): ("Comprehensive Dental Consultation", "Keng qamrovli stomatologik ko'rik", "Комплексная стоматологическая консультация"),
        
        (2, 1): ("Clinical Consultation & Initial Assessment", "Klinik konsultatsiya va birlamchi ko'rik", "Клиническая консультация и первичный осмотр"),
        (2, 2): ("Chest Pain & Cardiovascular Evaluation", "Ko'krak og'rig'i va yurak-qon tomir tekshiruvi", "Боль в груди и оценка сердечно-сосудистой системы"),
        (2, 3): ("Respiratory Symptoms & Dyspnea", "Nafas qisishi va o'pka simptomlari", "Одышка и респираторные симптомы"),
        (2, 4): ("Gastrointestinal & Peptic History", "Oshqozon-ichak va me'da yarasi anamnezi", "ЖКТ и анамнез язвенной болезни"),
        (2, 5): ("Neurological Assessment & Headaches", "Nevrologik tekshiruv va bosh og'rig'i", "Неврологический осмотр и головные боли"),
        (2, 6): ("Severe Infection & Sepsis Assessment", "Og'ir infeksiya va sepsis baholash", "Тяжелая инфекция и оценка сепсиса"),
        (2, 7): ("Acute Abdominal Pain Assessment", "O'tkir qorin og'rig'i anamnezi", "Острый живот и абдоминальная боль"),
        (2, 8): ("Urinary Tract Infection & Pyelonephritis", "Siydik yo'llari infeksiyasi va piyelonefrit", "Инфекция мочевыводящих путей и пиелонефрит"),
        (2, 9): ("Anemia & Fatigue Investigation", "Anemiya va surunkali holsizlik", "Анемия и хроническая усталость"),
        (2, 10): ("Metabolic & Endocrine Review", "Endokrin va metabolik ko'rik (Qandli diabet)", "Метаболический и эндокринный скрининг")
    }

    for spec in specs:
        spec_id = spec['id']
        spec_name = spec['name']
        spec_folder = spec['folder']

        for mod_num in range(1, 11):
            # Locate folder
            matched_dir = None
            for d in os.listdir(spec_folder):
                dp = os.path.join(spec_folder, d)
                if os.path.isdir(dp):
                    digits = re.findall(r'\d+', d)
                    if digits and int(digits[0]) == mod_num:
                        matched_dir = dp
                        break

            if not matched_dir:
                continue

            all_docs = [os.path.join(matched_dir, f) for f in os.listdir(matched_dir) if f.endswith('.docx') and not f.startswith('~$')]

            all_vocab = []
            all_phrases = []
            all_grammar = []
            all_dialogue = []

            for doc_path in all_docs:
                base_name = os.path.basename(doc_path).lower()
                p_lines, t_data = read_docx_clean(doc_path)

                if 'dialog' in base_name:
                    turns = parse_dialogue_file(p_lines)
                    if turns:
                        all_dialogue = turns

                if 'vocab' in base_name or 'phrase' in base_name:
                    v_items = parse_vocab_tables(t_data, p_lines)
                    p_items = parse_phrasebook_tables(t_data, p_lines)
                    if v_items:
                        all_vocab.extend(v_items)
                    if p_items:
                        all_phrases.extend(p_items)

                if 'grammar' in base_name:
                    g_items = parse_grammar_details(p_lines, t_data)
                    if g_items:
                        all_grammar.extend(g_items)

            # Deduplicate vocab by word
            unique_vocab = []
            seen_v = set()
            for v in all_vocab:
                if v['word'].lower() not in seen_v:
                    seen_v.add(v['word'].lower())
                    unique_vocab.append(v)

            # Deduplicate phrases by phrase
            unique_phrases = []
            seen_p = set()
            for p in all_phrases:
                if p['phrase'].lower() not in seen_p:
                    seen_p.add(p['phrase'].lower())
                    unique_phrases.append(p)

            if not all_grammar:
                all_grammar = parse_grammar_details([], [])

            title_en, title_uz, title_ru = titles_map.get((spec_id, mod_num), (f"Module {mod_num}", f"Modul {mod_num}", f"Модуль {mod_num}"))
            quizzes = generate_module_quizzes(title_en, unique_vocab, all_grammar)

            module_data = {
                'specialty_id': spec_id,
                'specialty_name': spec_name,
                'order_index': mod_num,
                'title': title_en,
                'title_uz': title_uz,
                'title_ru': title_ru,
                'title_en': title_en,
                'description': f"{spec_name} bo'yicha {title_uz} mavzusida klinik muloqot va tibbiy ingliz tili darsi.",
                'clinical_level': 'B2 Medical English',
                'target_role': 'Dental Clinician / Attending Physician',
                'scenario_prompt': f"You are a virtual patient presenting with symptoms related to {title_en}. Speak in natural, everyday patient English.",
                'patient_profile': {
                    'age': 30 + mod_num * 3,
                    'gender': 'Male' if mod_num % 2 == 1 else 'Female',
                    'chief_complaint': title_uz,
                    'history': f"Patient presents with complaints of {title_en.lower()}."
                },
                'grammar': all_grammar,
                'vocabulary': unique_vocab,
                'phrasebook': unique_phrases,
                'dialogue': all_dialogue,
                'tests': quizzes
            }

            # Save in module directory
            out_file = os.path.join(matched_dir, 'module_data.json')
            with open(out_file, 'w', encoding='utf-8') as f:
                json.dump(module_data, f, indent=2, ensure_ascii=False)

            print(f"[{spec_name}] Modul {mod_num}: {title_en} -> Vocab: {len(unique_vocab)}, Phrases: {len(unique_phrases)}, Grammar: {len(all_grammar)}, Dialogue: {len(all_dialogue)}")
            all_modules.append(module_data)

    # Master files in datas/
    stomatology_curriculum = [m for m in all_modules if m['specialty_id'] == 1]
    davolash_curriculum = [m for m in all_modules if m['specialty_id'] == 2]

    with open('datas/stomatology_curriculum.json', 'w', encoding='utf-8') as f:
        json.dump(stomatology_curriculum, f, indent=2, ensure_ascii=False)
    with open('datas/davolash_ishi_curriculum.json', 'w', encoding='utf-8') as f:
        json.dump(davolash_curriculum, f, indent=2, ensure_ascii=False)
    with open('datas/all_medical_curriculum.json', 'w', encoding='utf-8') as f:
        json.dump(all_modules, f, indent=2, ensure_ascii=False)

    print("\n✅ Barcha 20 ta modul JSON fayllari datas/ ichida muvaffaqiyatli saqlandi!")

if __name__ == '__main__':
    build_curriculum()
