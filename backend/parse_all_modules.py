import os
import sys
import json
import re
import docx

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

def get_text_and_tables(docx_path):
    doc = docx.Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        t_rows = []
        for row in table.rows:
            # deduplicate merged cells horizontally
            row_cells = []
            last_cell = None
            for cell in row.cells:
                text = cell.text.strip()
                if text != last_cell:
                    row_cells.append(text)
                    last_cell = text
            if any(row_cells):
                t_rows.append(row_cells)
        if t_rows:
            tables.append(t_rows)
    return paragraphs, tables

def inspect():
    print("--- Inspecting STOMOTOLOGY modules ---")
    stomatology_files = {}
    for root, dirs, files in os.walk('datas/STOMOTOLOGY'):
        for f in files:
            if f.endswith('.docx') and not f.startswith('~$'):
                p = os.path.join(root, f)
                pars, tabs = get_text_and_tables(p)
                stomatology_files[p] = {
                    'paragraphs_count': len(pars),
                    'tables_count': len(tabs),
                    'first_paragraphs': pars[:5],
                    'tables_preview': [{'rows': len(t), 'header': t[0] if t else []} for t in tabs]
                }
                print(f"File: {p} | Pars: {len(pars)} | Tables: {len(tabs)}")

    print("\n--- Inspecting DAVOLASH ISHI modules ---")
    davolash_files = {}
    for root, dirs, files in os.walk('datas/DAVOLASH ISHI'):
        for f in files:
            if f.endswith('.docx') and not f.startswith('~$'):
                p = os.path.join(root, f)
                pars, tabs = get_text_and_tables(p)
                davolash_files[p] = {
                    'paragraphs_count': len(pars),
                    'tables_count': len(tabs),
                    'first_paragraphs': pars[:5],
                    'tables_preview': [{'rows': len(t), 'header': t[0] if t else []} for t in tabs]
                }
                print(f"File: {p} | Pars: {len(pars)} | Tables: {len(tabs)}")

    with open('datas/inspection_summary.json', 'w', encoding='utf-8') as out:
        json.dump({
            'stomatology': stomatology_files,
            'davolash_ishi': davolash_files
        }, out, indent=2, ensure_ascii=False)
    print("\nWrote datas/inspection_summary.json")

if __name__ == '__main__':
    inspect()
